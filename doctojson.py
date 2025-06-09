#!/usr/bin/env python3
"""
Working Word to JSON Converter
Simple and reliable converter that works!
"""

try:
    from docx import Document
    import sys
    import importlib
    # Force import the built-in json module
    json_module = importlib.import_module('json')
    import re
    from pathlib import Path
    print("✅ All modules imported successfully!")
except ImportError as e:
    print(f"❌ Missing module: {e}")
    print("Please install: pip install python-docx")
    exit(1)

def convert_word_to_json(file_path, output_dir):
    """Convert Word document to JSON format"""
    print(f"🔄 Converting: {file_path}")
    print(f"📂 Output directory: {output_dir}")
    
    try:
        # Check if file exists
        if not Path(file_path).exists():
            print(f"❌ File not found: {file_path}")
            return False
        
        # Load document
        doc = Document(file_path)
        print("📖 Document loaded successfully!")
        
        # Initialize result
        doc_name = Path(file_path).stem
        result = {
            "document_title": doc_name.replace('_', ' '),
            "content": []
        }
        
        current_section = None
        media_counter = 1
        
        # Process each paragraph
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            if not text:
                continue
                
            # Skip table of contents
            if any(skip in text.lower() for skip in ['table of content', 'contents', '....', '___']):
                continue
            
            print(f"📝 Processing paragraph {i+1}: {text[:50]}...")
            
            # Check formatting
            is_bold = False
            is_large = False
            
            for run in para.runs:
                if run.bold:
                    is_bold = True
                if run.font.size and run.font.size.pt >= 14:
                    is_large = True
            
            # Check style
            is_heading = 'heading' in para.style.name.lower()
            
            # Determine content type
            text_lower = text.lower()
            
            if is_heading or (is_bold and is_large):
                # This is a section
                content_type = "section"
                current_section = {
                    "type": "section",
                    "title": text,
                    "content": []
                }
                result["content"].append(current_section)
                print(f"  📁 Section: {text}")
                
            elif is_bold or text_lower.startswith(('step', 'access', 'enable', 'change', 'update')) or re.match(r'^\d+\.', text):
                # This is a step
                content_type = "step"
                content_item = {
                    "type": "step",
                    "text": text
                }
                
                if current_section:
                    current_section["content"].append(content_item)
                else:
                    result["content"].append(content_item)
                print(f"  🔢 Step: {text[:30]}...")
                
            elif any(media_word in text_lower for media_word in ['image', 'media', '.png', '.gif', 'figure']):
                # This is media
                media_path = f"https://github.com/SavvyGaikwad/img/blob/main/{doc_name.replace('_', ' ')}/{media_counter}.png"
                if '.gif' in text_lower:
                    media_path = media_path.replace('.png', '.gif')
                
                media_item = {
                    "type": "media",
                    "path": media_path
                }
                
                if current_section:
                    current_section["content"].append(media_item)
                else:
                    result["content"].append(media_item)
                    
                media_counter += 1
                print(f"  🖼️ Media: {media_path}")
                
            else:
                # This is info
                content_item = {
                    "type": "info",
                    "text": text
                }
                
                if current_section:
                    current_section["content"].append(content_item)
                else:
                    result["content"].append(content_item)
                print(f"  ℹ️ Info: {text[:30]}...")
        
        # Process tables
        if doc.tables:
            print(f"📊 Processing {len(doc.tables)} tables...")
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    table_item = {
                        "type": "info",
                        "text": "\n".join(table_text)
                    }
                    
                    if current_section:
                        current_section["content"].append(table_item)
                    else:
                        result["content"].append(table_item)
        
        # Save to JSON in specified output directory
        output_filename = Path(file_path).stem + '.json'
        output_path = Path(output_dir) / output_filename
        
        # Create output directory if it doesn't exist
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json_module.dump(result, f, indent=2, ensure_ascii=False)
        
        print(f"✅ SUCCESS! JSON saved to: {output_path}")
        
        # Show summary
        sections = len([item for item in result['content'] if item.get('type') == 'section'])
        steps = len([item for item in result['content'] if item.get('type') == 'step'])
        info_items = len([item for item in result['content'] if item.get('type') == 'info'])
        media_items = len([item for item in result['content'] if item.get('type') == 'media'])
        
        print(f"📊 SUMMARY:")
        print(f"   📁 Sections: {sections}")
        print(f"   🔢 Steps: {steps}")
        print(f"   ℹ️ Info items: {info_items}")
        print(f"   🖼️ Media items: {media_items}")
        
        return True
        
    except Exception as e:
        print(f"❌ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Main function"""
    print("🚀 Word to JSON Converter Starting...")
    print("=" * 50)
    
    # Input Word document path
    file_path = r"C:\Users\hp\Downloads\English-20250601T110944Z-1-001\user manual\Modules\Work Requests\Requests Settings.docx"
    
    # Output JSON directory
    output_dir = r"C:\Users\hp\Downloads\final"
    
    success = convert_word_to_json(file_path, output_dir)
    
    print("=" * 50)
    if success:
        print("🎉 CONVERSION COMPLETED SUCCESSFULLY!")
    else:
        print("💥 CONVERSION FAILED!")
    
    print("Press Enter to exit...")
    input()

if __name__ == "__main__":
    main()